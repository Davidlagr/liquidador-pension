import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from datetime import date, datetime, timedelta
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from data_processor import extraer_tabla_cruda, limpiar_y_estandarizar, aplicar_regla_simultaneidad
from logic import LiquidadorPension

# ==========================================
# 1. CONFIGURACIÓN DE PÁGINA
# ==========================================
st.set_page_config(page_title="Liquidador Pensional Pro", layout="wide", page_icon="⚖️")

# ==========================================
# 2. SISTEMA DE AUTENTICACIÓN
# ==========================================
if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False

if not st.session_state.autenticado:
    st.markdown("<h2 style='text-align: center;'>🔒 Acceso al Sistema Pensional</h2>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        with st.form("login_form"):
            usuario = st.text_input("Usuario")
            clave = st.text_input("Contraseña", type="password")
            submit = st.form_submit_button("Ingresar", use_container_width=True)
            
            if submit:
                if usuario == "80799289" and clave == "1808DV":
                    st.session_state.autenticado = True
                    st.rerun()
                else:
                    st.error("❌ Credenciales incorrectas. Verifica tu usuario y contraseña.")
    st.stop()

# --- CSS ---
st.markdown("""
    <style>
    .info-box { background-color: #e8f6f3; padding: 15px; border-radius: 8px; border-left: 5px solid #1abc9c; margin-bottom: 15px; }
    .ibl-box { background-color: #f4f6f7; padding: 10px; border-radius: 5px; border: 1px solid #d5dbdb; text-align: center; }
    .status-ok { color: #27ae60; font-weight: bold; }
    .status-alert { color: #d35400; font-weight: bold; }
    .formula-box { background-color: #fef9e7; padding: 15px; border-radius: 8px; border-left: 5px solid #f1c40f; margin-top: 10px; }
    </style>
""", unsafe_allow_html=True)

st.title("⚖️ Liquidador Pensional: Análisis Técnico & Jurídico")

# --- ESTADO Y VARIABLES GLOBALES ---
if 'df_crudo' not in st.session_state: st.session_state.df_crudo = None
if 'df_final' not in st.session_state: st.session_state.df_final = None

SMLMV_ACTUAL_2026 = 1750905.0

# ==========================================
# FUNCIONES AUXILIARES JURÍDICAS Y DE LIMPIEZA
# ==========================================
def limpiar_valor_ibc_robusto(val):
    if pd.isna(val):
        return 0.0
    
    val_str = str(val).strip().replace('$', '').replace('COP', '').replace(' ', '')
    if not val_str or val_str.lower() in ['nan', 'none']:
        return 0.0

    if ',' in val_str and '.' in val_str:
        val_str = val_str.replace('.', '').replace(',', '.')
    elif '.' in val_str:
        partes = val_str.split('.')
        if len(partes) > 1 and len(partes[-1]) == 3:
            val_str = ''.join(partes)
    elif ',' in val_str:
        partes = val_str.split(',')
        if len(partes) > 1 and len(partes[-1]) == 3:
            val_str = ''.join(partes)
        else:
            val_str = val_str.replace(',', '.')

    try:
        return float(val_str)
    except ValueError:
        return 0.0

def get_requisitos_estatus(genero, fecha_estatus, fecha_cumple_edad=None):
    edad_req = 62 if genero == "Masculino" else 57
    
    if genero == "Masculino":
        semanas_req = 1300
        nota = "Aplica regla general Ley 797/2003 (1300 semanas)."
    else:
        if pd.isna(fecha_estatus):
            anio_actual = datetime.now().year
            anio_cumple = fecha_cumple_edad.year if fecha_cumple_edad else anio_actual
            
            anio_proy = max(anio_actual, anio_cumple)
            
            if anio_proy < 2026:
                anio_proy = 2026 
                
            if anio_proy == 2026:
                semanas_req = 1250
            else:
                descenso = 50 + ((anio_proy - 2026) * 25)
                semanas_req = max(1000, 1300 - descenso)
                
            nota = f"Aún no consolida estatus. Se proyecta exigencia de {semanas_req} semanas para el año {anio_proy} acorde a Sentencia C-197/23."
        else:
            anio = fecha_estatus.year
            if anio < 2026:
                semanas_req = 1300
                nota = f"Consolidó estatus en {anio} (antes de 2026). Aplica exigencia de 1300 semanas."
            else:
                if anio == 2026:
                    semanas_req = 1250
                else:
                    descenso = 50 + ((anio - 2026) * 25)
                    semanas_req = max(1000, 1300 - descenso)
                nota = f"Consolidó estatus en {anio}. Aplica disminución progresiva constitucional: {semanas_req} semanas."

    return edad_req, semanas_req, nota

def desglosar_formula_tasa(ibl, semanas_totales, semanas_req, smlmv_ref):
    s = ibl / smlmv_ref if smlmv_ref > 0 else 1
    tasa_base = 65.5 - (0.5 * s)
    tasa_base = max(55.0, min(tasa_base, 65.5))
    
    semanas_adicionales = max(0, semanas_totales - semanas_req)
    bloques_50 = int(semanas_adicionales // 50)
    incremento = bloques_50 * 1.5
    
    tasa_final = min(80.0, tasa_base + incremento)
    
    return {
        "s": s, "tasa_base": tasa_base, "semanas_adicionales": semanas_adicionales,
        "bloques": bloques_50, "incremento": incremento, "tasa_final": tasa_final
    }

# ==========================================
# GENERADOR DE REPORTE WORD
# ==========================================
def generar_reporte_completo(perfil, fechas, liq_data, req_data, proyeccion=None):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    import matplotlib.ticker as mtick
    fmt = '${x:,.0f}'
    tick = mtick.StrMethodFormatter(fmt)

    tit = doc.add_heading('DICTAMEN TÉCNICO PENSIONAL', 0)
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha de Emisión: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("_" * 70)

    doc.add_heading('1. MARCO NORMATIVO Y METODOLÓGICO', level=1)
    doc.add_paragraph(
        "El presente análisis técnico-jurídico se fundamenta en los parámetros de la Ley 100 de 1993 y las "
        "modificaciones introducidas por la Ley 797 de 2003.\n\n"
        "• Cálculo IBL: Indexación de los IBC históricos utilizando el IPC del DANE, comparando "
        "los últimos 10 años cotizados frente a toda la vida laboral (Principio de Favorabilidad).\n\n"
        "• Reducción de Semanas (Mujeres): En acatamiento a la jurisprudencia constitucional (Sentencia C-197/23), "
        "se aplica la disminución progresiva del número de semanas exigidas para las mujeres que consoliden "
        "su estatus pensional a partir del 1° de enero de 2026."
    )

    doc.add_heading('2. REQUISITOS PARA EL ESTATUS', level=1)
    t_req = doc.add_table(rows=1, cols=2)
    t_req.style = 'Table Grid'
    for k, v in [("Edad Exigida", f"{req_data['edad']} años"), ("Semanas Exigidas", f"{req_data['semanas']}"), ("Fundamento", req_data['nota'])]:
        r = t_req.add_row().cells
        r[0].text, r[1].text = k, str(v)

    doc.add_heading('3. CONSOLIDACIÓN DEL DERECHO', level=1)
    t_est = doc.add_table(rows=1, cols=2)
    t_est.style = 'Table Grid'
    for k, v in [
        ("Afiliado", perfil['nombre']),
        ("Cumplimiento Edad", fechas['fecha_cumple_edad'].strftime('%d/%m/%Y')),
        ("Cumplimiento Semanas", fechas['fecha_cumple_semanas'].strftime('%d/%m/%Y') if fechas['fecha_cumple_semanas'] else "No cumplido"),
        ("FECHA ESTATUS PENSIONAL", fechas['fecha_estatus'].strftime('%d/%m/%Y') if fechas['tiene_estatus'] else "NO ADQUIRIDO"),
        ("FECHA CORTE (INDEXACIÓN)", fechas['fecha_corte'].strftime('%d/%m/%Y')),
    ]:
        r = t_est.add_row().cells
        r[0].text, r[1].text = k, str(v)

    doc.add_page_break()
    doc.add_heading('4. RESULTADO DE LA LIQUIDACIÓN ACTUAL', level=1)
    
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Light Shading Accent 1'
    for k, v in [
        ("Semanas Totales", f"{liq_data['semanas']:,.2f}"),
        ("IBL 10 Años", f"${liq_data['ibl_10']:,.0f}"),
        ("IBL Toda la Vida", f"${liq_data['ibl_vida']:,.0f}"),
        ("IBL APLICADO", f"${liq_data['ibl']:,.0f} ({liq_data['origen_ibl']})"),
        ("Tasa Reemplazo Aplicada", f"{liq_data['tasa']:.2f}%"),
        ("MESADA PENSIONAL", f"${liq_data['mesada']:,.0f}")
    ]:
        r = t2.add_row().cells
        r[0].text, r[1].text = k, v

    fig_ibl, ax_ibl = plt.subplots(figsize=(5, 3))
    ax_ibl.bar(["Últimos 10", "Toda Vida"], [liq_data['ibl_10'], liq_data['ibl_vida']], color=['#3498db', '#2ecc71'])
    ax_ibl.set_title("Comparativo Ingreso Base de Liquidación (IBL)")
    ax_ibl.yaxis.set_major_formatter(tick)
    mem_ibl = BytesIO()
    fig_ibl.savefig(mem_ibl, format='png', bbox_inches='tight')
    doc.add_paragraph("\n")
    doc.add_picture(mem_ibl, width=Inches(4.5))
    mem_ibl.close()
    plt.close(fig_ibl)

    doc.add_heading('4.1 FÓRMULA DE TASA DE REEMPLAZO', level=2)
    f_data = liq_data['formula_tasa']
    doc.add_paragraph(
        f"1. Cálculo SMLMV del IBL (s): {f_data['s']:.2f} salarios mínimos.\n"
        f"2. Tasa Base [65.5 - (0.5 * s)]: {f_data['tasa_base']:.2f}%\n"
        f"3. Semanas Adicionales ({liq_data['semanas']:,.2f} - {req_data['semanas']}): {f_data['semanas_adicionales']:.2f}\n"
        f"4. Bloques de 50 semanas: {f_data['bloques']}\n"
        f"5. Incremento en Tasa ({f_data['bloques']} x 1.5%): +{f_data['incremento']:.2f}%\n"
        f"6. TASA FINAL (Tope Legal 80%): {f_data['tasa_final']:.2f}%"
    )

    def agregar_tabla_soporte(df_sop):
        if not df_sop.empty:
            t = doc.add_table(rows=1, cols=3)
            t.style = 'Table Grid'
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text = 'Periodo', 'IBC Histórico', 'IBC Actualizado'
            
            filas_mostrar = pd.concat([df_sop.head(30), df_sop.tail(10)]) if len(df_sop) > 40 else df_sop
            for _, row in filas_mostrar.iterrows():
                rc = t.add_row().cells
                rc[0].text = f"{row['Desde'].strftime('%m/%Y')} - {row['Hasta'].strftime('%m/%Y')}"
                rc[1].text, rc[2].text = f"${row['IBC_Historico']:,.0f}", f"${row['IBC_Actualizado']:,.0f}"

    doc.add_page_break()
    doc.add_heading('ANEXO 1: DETALLE ÚLTIMOS 10 AÑOS (ACTUAL)', level=1)
    agregar_tabla_soporte(liq_data['df_soporte_10'])

    doc.add_heading('ANEXO 2: DETALLE TODA LA VIDA (ACTUAL)', level=1)
    agregar_tabla_soporte(liq_data['df_soporte_vida'])

    if proyeccion:
        doc.add_page_break()
        doc.add_heading('5. PROYECCIÓN ESTRATÉGICA DE MEJORA PENSIONAL', level=1)
        doc.add_paragraph("Esquema de viabilidad financiera liquidando salud (12.5%) y pensión (16%) sobre el 100% del IBC proyectado. "
                          "El sistema proyecta el SMLMV futuro aplicando la tasa de aumento anual estimada mediante interés compuesto.")
        
        t3 = doc.add_table(rows=1, cols=2)
        t3.style = 'Table Grid'
        for k, v in [
            ("Perfil de Cotizante", proyeccion['estrategia']),
            ("Años Proyectados", f"{proyeccion['anios']} años"),
            ("INVERSIÓN TOTAL DE BOLSILLO", f"${proyeccion['inversion']:,.0f}"),
            ("NUEVA MESADA PROYECTADA", f"${proyeccion['mesada_fut']:,.0f}"),
            ("Incremento de Mesada (Delta)", f"${proyeccion['delta']:,.0f} mensuales"),
            ("Tiempo de Retorno (ROI)", f"{proyeccion['roi']:.1f} Años tras pensionarse")
        ]:
            r = t3.add_row().cells
            r[0].text, r[1].text = k, str(v)

        fig_p1, ax_p1 = plt.subplots(figsize=(5, 3))
        ax_p1.bar(["Actual", "Proyectada"], [liq_data['mesada'], proyeccion['mesada_fut']], color=['#e74c3c', '#27ae60'])
        ax_p1.set_title("Incremento de Mesada Pensional")
        ax_p1.yaxis.set_major_formatter(tick)
        mem_p1 = BytesIO()
        fig_p1.savefig(mem_p1, format='png', bbox_inches='tight')
        doc.add_paragraph("\n")
        doc.add_picture(mem_p1, width=Inches(4.5))
        mem_p1.close()
        plt.close(fig_p1)
        
        # --- NUEVA SECCIÓN DE DESGLOSE FINANCIERO Y GRÁFICAS ---
        doc.add_heading('5.1 DESGLOSE DE COSTOS Y PROYECCIÓN DE APORTES', level=2)
        doc.add_paragraph("A continuación, se detalla la evolución del Ingreso Base de Cotización (IBC) proyectado y el costo de los aportes (Salud + Pensión), indexados con la inflación estimada.")
        
        df_inv = pd.DataFrame(proyeccion['detalle_inversion'])
        t4 = doc.add_table(rows=1, cols=5)
        t4.style = 'Table Grid'
        hdr_cells = t4.rows[0].cells
        hdr_cells[0].text = 'Año Proy.'
        hdr_cells[1].text = 'SMLMV Est.'
        hdr_cells[2].text = 'IBC Mensual'
        hdr_cells[3].text = 'Aporte Mensual'
        hdr_cells[4].text = 'Costo Anual'
        
        for _, row in df_inv.iterrows():
            row_cells = t4.add_row().cells
            row_cells[0].text = str(int(row['Año']))
            row_cells[1].text = f"${row['SMLMV Proyectado']:,.0f}"
            row_cells[2].text = f"${row['IBC Mes']:,.0f}"
            row_cells[3].text = f"${row['Costo Mes']:,.0f}"
            row_cells[4].text = f"${row['Costo Anual']:,.0f}"

        # Gráfico adicional para la proyección de costos
        fig_p2, ax_p2 = plt.subplots(figsize=(6, 3))
        ax_p2.plot(df_inv['Año'], df_inv['Costo Mes'], marker='o', linestyle='-', color='#8e44ad')
        ax_p2.set_title("Evolución del Costo Mensual de Aportes")
        ax_p2.set_xlabel("Año de Proyección")
        ax_p2.set_ylabel("Costo Mensual ($)")
        ax_p2.yaxis.set_major_formatter(tick)
        ax_p2.set_xticks(df_inv['Año'])
        mem_p2 = BytesIO()
        fig_p2.savefig(mem_p2, format='png', bbox_inches='tight')
        doc.add_paragraph("\n")
        doc.add_picture(mem_p2, width=Inches(5.0))
        mem_p2.close()
        plt.close(fig_p2)
            
        doc.add_page_break()
        doc.add_heading('5.2 SOPORTES: CÁLCULO DE LA MESADA FUTURA', level=2)
        doc.add_paragraph(f"Escenario más favorable aplicado para el futuro: {proyeccion['origen_ibl_fut']}.")
        
        doc.add_heading('DETALLE PROYECCIÓN: ÚLTIMOS 10 AÑOS', level=3)
        agregar_tabla_soporte(proyeccion['det_10_fut'])

        doc.add_heading('DETALLE PROYECCIÓN: TODA LA VIDA', level=3)
        agregar_tabla_soporte(proyeccion['det_vida_fut'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# INTERFAZ (SIDEBAR)
# ==========================================
with st.sidebar:
    st.header("👤 Datos")
    nombre = st.text_input("Nombre", "Usuario")
    genero = st.radio("Género", ["Masculino", "Femenino"])
    fecha_nac = st.date_input("Nacimiento", value=date(1975, 1, 1), min_value=date(1900,1,1), max_value=datetime.now().date())
    
    st.divider()
    aplicar_tope = st.checkbox("Tope 1800 Semanas", value=True)
    
    if st.button("🚪 Cerrar Sesión"):
        st.session_state.autenticado = False
        st.session_state.df_crudo = None
        st.session_state.df_final = None
        st.rerun()
        
    if st.button("🔄 Reiniciar App"):
        st.session_state.df_crudo = None
        st.session_state.df_final = None
        st.rerun()

# ==========================================
# LÓGICA PRINCIPAL
# ==========================================
if st.session_state.df_final is None:
    st.info("📂 Carga el PDF de Historia Laboral")
    uploaded_file = st.file_uploader("Archivo PDF", type="pdf")

    if uploaded_file:
        if st.session_state.df_crudo is None:
            st.session_state.df_crudo = extraer_tabla_cruda(uploaded_file)
        
        df = st.session_state.df_crudo
        if df is not None and not df.empty:
            st.dataframe(df.head(3))
            cols = df.columns.tolist()
            c1, c2, c3, c4 = st.columns(4)
            cd = c1.selectbox("Desde", cols, index=2 if len(cols)>2 else 0)
            ch = c2.selectbox("Hasta", cols, index=3 if len(cols)>3 else 0)
            ci = c3.selectbox("IBC", cols, index=4 if len(cols)>4 else 0)
            cs = c4.selectbox("Semanas", cols, index=len(cols)-1)
            
            if st.button("Procesar"):
                df_procesar = df.copy()
                df_procesar[ci] = df_procesar[ci].apply(limpiar_valor_ibc_robusto)
                
                clean = limpiar_y_estandarizar(df_procesar, cd, ch, ci, cs)
                if not clean.empty:
                    st.session_state.df_final = aplicar_regla_simultaneidad(clean)
                    st.rerun()
                else:
                    st.error("Error validando las columnas.")

else:
    df = st.session_state.df_final
    liq = LiquidadorPension(df, genero, fecha_nac)
    
    fechas_clave = liq.determinar_fechas_clave()
    ibl_10, det_10 = liq.calcular_ibl_indexado(fechas_clave['fecha_corte'], "ultimos_10")
    ibl_vida, det_vida = liq.calcular_ibl_indexado(fechas_clave['fecha_corte'], "toda_vida")
    
    ibl_def = max(ibl_10, ibl_vida)
    origen_ibl = "Últimos 10 Años" if ibl_10 >= ibl_vida else "Toda la Vida"
    
    total_sem = df['Semanas'].sum()
    
    edad_req, semanas_req, nota_req = get_requisitos_estatus(genero, fechas_clave['fecha_estatus'], fechas_clave['fecha_cumple_edad'])
    formula_tasa = desglosar_formula_tasa(ibl_def, total_sem, semanas_req, SMLMV_ACTUAL_2026)
    
    mesada, tasa, info = liq.calcular_tasa_reemplazo_797(
        ibl_def, total_sem, datetime.now().year, aplicar_tope
    )

    tab1, tab2 = st.tabs(["📊 DIAGNÓSTICO JURÍDICO", "📈 PROYECCIÓN DE MEJORA"])
    
    with tab1:
        st.subheader(f"Dictamen de Estatus: {nombre}")
        
        st.markdown(f"""
        <div style='background-color: #f5eef8; padding: 10px; border-radius: 5px; border-left: 4px solid #8e44ad; margin-bottom: 15px;'>
            <b>📌 Requisitos Normativos para el Afiliado:</b><br>
            Edad Exigida: <b>{edad_req} años</b> | Semanas Exigidas: <b>{semanas_req}</b><br>
            <i>{nota_req}</i>
        </div>
        """, unsafe_allow_html=True)
        
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            estado_texto = 'PENDIENTE (REQUISITOS NO CUMPLIDOS)'
            if fechas_clave['tiene_estatus']:
                estado_texto = f"{fechas_clave['fecha_estatus'].strftime('%d/%m/%Y')} ✅"
                
            st.markdown(f"""
            <div class='info-box'>
                <b>Fecha Cumplimiento Edad:</b> {fechas_clave['fecha_cumple_edad'].strftime('%d/%m/%Y')}<br>
                <b>Fecha Cumplimiento Semanas:</b> {fechas_clave['fecha_cumple_semanas'].strftime('%d/%m/%Y') if fechas_clave['fecha_cumple_semanas'] else 'No cumplido'}<br>
                <b>FECHA DE ESTATUS:</b> {estado_texto}
            </div>
            """, unsafe_allow_html=True)
            
        with col_f2:
            st.markdown(f"""
            <div class='info-box' style='border-color: #e67e22; background-color: #fcf3cf;'>
                <b>FECHA DE CORTE (INDEXACIÓN):</b> {fechas_clave['fecha_corte'].strftime('%d/%m/%Y')}<br>
                <b>Razón:</b> {fechas_clave['razon_corte']}<br>
                <b>FECHA EFECTIVIDAD:</b> {fechas_clave['fecha_efectividad'].strftime('%d/%m/%Y')}
            </div>
            """, unsafe_allow_html=True)
            
        st.divider()
        
        c1, c2, c3 = st.columns(3)
        c1.metric("Semanas Totales", f"{total_sem:,.2f}")
        c2.metric("Mesada Pensional", f"${mesada:,.0f}")
        c3.metric("Tasa Reemplazo Aplicada", f"{tasa:.2f}%")

        with st.expander("📐 Ver fórmula de cálculo de la Tasa de Reemplazo"):
            st.markdown(f"""
            <div class='formula-box'>
                <strong>1. Cálculo Base (Art. 34 Ley 100):</strong><br>
                • SMLMV del IBL (s): {formula_tasa['s']:.2f} salarios<br>
                • Tasa Base [65.5 - (0.5 * s)]: <b>{formula_tasa['tasa_base']:.2f}%</b><br><br>
                <strong>2. Incremento por Semanas Adicionales:</strong><br>
                • Semanas Adicionales ({total_sem:,.2f} totales - {semanas_req} exigidas): {formula_tasa['semanas_adicionales']:.2f}<br>
                • Bloques de 50 semanas: {formula_tasa['bloques']}<br>
                • Incremento ({formula_tasa['bloques']} x 1.5%): <b>+{formula_tasa['incremento']:.2f}%</b><br><br>
                <strong>3. Tasa Final (Tope 80%): <b>{formula_tasa['tasa_final']:.2f}%</b></strong>
            </div>
            """, unsafe_allow_html=True)

        st.divider()

        st.markdown("#### 🆚 Análisis Comparativo de Ingreso Base (IBL)")
        col_ibl_L, col_ibl_R = st.columns(2)
        with col_ibl_L:
            st.markdown(f"""<div class='ibl-box'><h4>Últimos 10 Años</h4><h2>${ibl_10:,.0f}</h2></div>""", unsafe_allow_html=True)
        with col_ibl_R:
             st.markdown(f"""<div class='ibl-box'><h4>Toda la Vida</h4><h2>${ibl_vida:,.0f}</h2></div>""", unsafe_allow_html=True)
            
        st.caption(f"El sistema aplicó automáticamente: **{origen_ibl}** por ser el escenario más favorable para el afiliado.")
        chart_data = pd.DataFrame({"Monto": [ibl_10, ibl_vida]}, index=["Últimos 10 Años", "Toda la Vida"])
        st.bar_chart(chart_data, color="#2E86C1")

        st.markdown("#### 📄 Soportes Técnicos Detallados")
        col_det_1, col_det_2 = st.columns(2)
        with col_det_1:
            with st.expander("🔍 Ver Detalle Últimos 10 Años (Actual)"):
                st.dataframe(det_10.style.format({'IBC_Historico': "${:,.0f}", 'IBC_Actualizado': "${:,.0f}", 'Factor_IPC': "{:.4f}"}))
        with col_det_2:
            with st.expander("🌍 Ver Detalle Toda la Vida (Actual)"):
                st.dataframe(det_vida.style.format({'IBC_Historico': "${:,.0f}", 'IBC_Actualizado': "${:,.0f}", 'Factor_IPC': "{:.4f}"}))

    with tab2:
        st.subheader("Simulación Financiera de Mejora Pensional")
        c_conf, c_res = st.columns([1, 2])
        
        with c_conf:
            st.markdown("### Parámetros de Inversión")
            opcion = st.radio("Perfil de Cotizante", ["Cotizante Independiente", "Dependiente + Extra Independiente"])
            
            smlmv_actual_proy = st.number_input("SMLMV Año Actual Base", value=SMLMV_ACTUAL_2026, step=100000.0)
            
            if opcion == "Cotizante Independiente":
                st.info("💡 **Independiente:** Inversión de Salud (12.5%) y Pensión (16%) sobre el **100% del IBC deseado** (Sin presunción del 40%).")
                smlmv_deseados = st.number_input("IBC Deseado (En SMLMV)", min_value=1.0, max_value=25.0, value=5.0, step=0.5)
                estrategia_texto = f"Independiente (IBC: {smlmv_deseados} SMLMV)"
                
                ibc_ref_ano_1 = smlmv_actual_proy * smlmv_deseados
                costo_mes_ano_1 = ibc_ref_ano_1 * 0.285
                st.success(f"**Ref. Año 1:**\nIBC Total: ${ibc_ref_ano_1:,.0f}\nCosto Mensual: ${costo_mes_ano_1:,.0f}")
            
            else:
                st.info("💡 **Dependiente:** Inversión del 28.5% se aplica **solo sobre el IBC extra** (La base actual la asume el empleador).")
                ultimo_ibc = float(df['IBC'].iloc[-1]) if not df.empty else smlmv_actual_proy
                st.write(f"**IBC Base (Empleador):** ${ultimo_ibc:,.0f}")
                
                smlmv_extra = st.number_input("IBC Extra a aportar (En SMLMV)", min_value=1.0, max_value=25.0, value=2.0, step=0.5)
                estrategia_texto = f"Dependiente + Extra Indep. ({smlmv_extra} SMLMV)"
                
                ibc_extra_ano_1 = smlmv_actual_proy * smlmv_extra
                costo_mes_ano_1 = ibc_extra_ano_1 * 0.285
                st.success(f"**Ref. Año 1:**\nIBC Extra (Inv): ${ibc_extra_ano_1:,.0f}\nCosto Extra Mensual: ${costo_mes_ano_1:,.0f}")

            anios_proy = st.slider("Años a realizar el aporte", 1, 15, 5)
            incremento_anual_smlmv = st.number_input("Est. Aumento Anual SMLMV (%)", value=5.0, step=0.5) / 100.0
            
            st.caption("📈 *Proyección basada en interés compuesto para simular la inflación del SMLMV y el ajuste real de aportes año a año.*")

        with c_res:
            filas_fut = []
            cur = df['Hasta'].max() + timedelta(days=1)
            inversion_total = 0
            detalle_inversion = []
            
            for m in range(anios_proy * 12):
                year_offset = m // 12
                smlmv_periodo = smlmv_actual_proy * ((1 + incremento_anual_smlmv) ** year_offset)
                
                if opcion == "Cotizante Independiente":
                    ibc_periodo = smlmv_deseados * smlmv_periodo
                    costo_mes = ibc_periodo * 0.285
                else:
                    ibc_dependiente_periodo = ultimo_ibc * ((1 + incremento_anual_smlmv) ** year_offset)
                    ibc_extra_periodo = smlmv_extra * smlmv_periodo
                    ibc_periodo = ibc_dependiente_periodo + ibc_extra_periodo
                    costo_mes = ibc_extra_periodo * 0.285
                
                inversion_total += costo_mes
                
                # Guardar el desglose anual en el mes 1 de cada año de proyección
                if m % 12 == 0:
                    detalle_inversion.append({
                        "Año": year_offset + 1,
                        "SMLMV Proyectado": smlmv_periodo,
                        "IBC Mes": ibc_periodo,
                        "Costo Mes": costo_mes,
                        "Costo Anual": costo_mes * 12
                    })

                filas_fut.append({
                    "Desde": cur, "Hasta": cur + timedelta(days=30), 
                    "IBC": ibc_periodo, "Semanas": 4.29
                })
                cur += timedelta(days=31)
            
            df_fut = pd.concat([df, pd.DataFrame(filas_fut)], ignore_index=True)
            liq_f = LiquidadorPension(df_fut, genero, fecha_nac)
            
            fechas_fut = liq_f.determinar_fechas_clave()
            
            ibl_10_fut, det_10_fut = liq_f.calcular_ibl_indexado(fechas_fut['fecha_corte'], "ultimos_10")
            ibl_vida_fut, det_vida_fut = liq_f.calcular_ibl_indexado(fechas_fut['fecha_corte'], "toda_vida")
            
            ibl_f = max(ibl_10_fut, ibl_vida_fut)
            origen_ibl_fut = "Últimos 10 Años" if ibl_10_fut >= ibl_vida_fut else "Toda la Vida"
            
            mes_f, tasa_f, _ = liq_f.calcular_tasa_reemplazo_797(ibl_f, df_fut['Semanas'].sum(), datetime.now().year + anios_proy, aplicar_tope)
            
            delta = mes_f - mesada
            roi = (inversion_total / (delta * 12)) if delta > 0 else 0
            
            st.markdown("### Resultados de la Proyección")
            m1, m2, m3 = st.columns(3)
            m1.metric("Nueva Mesada", f"${mes_f:,.0f}", f"+ ${delta:,.0f} mes")
            m2.metric("Inversión de Bolsillo", f"${inversion_total:,.0f}")
            if roi > 0: 
                m3.metric("Retorno (ROI)", f"{roi:.1f} años")
            else: 
                m3.error("Sin mejora")
                
            st.caption(f"**Escenario aplicado para liquidación futura:** {origen_ibl_fut}")
            
            chart_proy = pd.DataFrame({"Mesada": [mesada, mes_f]}, index=["Situación Actual", "Con Proyección"])
            st.bar_chart(chart_proy, color="#27AE60")

            # --- NUEVO EXPANSOR EN LA INTERFAZ DE USUARIO ---
            with st.expander("💸 Ver Desglose Anual de Costos de Inversión"):
                df_inv_ui = pd.DataFrame(detalle_inversion)
                st.dataframe(df_inv_ui.style.format({
                    'SMLMV Proyectado': '${:,.0f}',
                    'IBC Mes': '${:,.0f}',
                    'Costo Mes': '${:,.0f}',
                    'Costo Anual': '${:,.0f}'
                }))

            t_d1, t_d2 = st.tabs(["📑 Detalle 10 Años (Proyectado)", "📑 Detalle Toda la Vida (Proyectado)"])
            with t_d1:
                st.dataframe(det_10_fut.style.format({'IBC_Historico': "${:,.0f}", 'IBC_Actualizado': "${:,.0f}"}))
            with t_d2:
                st.dataframe(det_vida_fut.style.format({'IBC_Historico': "${:,.0f}", 'IBC_Actualizado': "${:,.0f}"}))
            
            proyeccion_data = {
                "estrategia": estrategia_texto, "anios": anios_proy, "inversion": inversion_total, 
                "mesada_fut": mes_f, "delta": delta, "roi": roi,
                "origen_ibl_fut": origen_ibl_fut, "det_10_fut": det_10_fut, "det_vida_fut": det_vida_fut,
                "detalle_inversion": detalle_inversion
            }

    # --- BOTÓN WORD Y APROBACIÓN ---
    st.sidebar.markdown("---")
    st.sidebar.subheader("📄 Generación de Informe")
    
    incluir_proyeccion = st.sidebar.checkbox("✅ Aprobar e Incluir Proyección al Dictamen", value=False)
    
    liq_data = {
        "semanas": total_sem, "ibl": ibl_def, "origen_ibl": origen_ibl, 
        "tasa": tasa, "mesada": mesada, "ibl_10": ibl_10, "ibl_vida": ibl_vida,
        "formula_tasa": formula_tasa,
        "df_soporte_10": det_10, "df_soporte_vida": det_vida 
    }
    perfil = {"nombre": nombre, "fecha_nac": fecha_nac.strftime('%d/%m/%Y')}
    req_data = {"edad": edad_req, "semanas": semanas_req, "nota": nota_req}
    
    docx = generar_reporte_completo(perfil, fechas_clave, liq_data, req_data, proyeccion_data if incluir_proyeccion else None)
    
    st.sidebar.download_button(
        label="📥 Descargar Dictamen Técnico (Word)", 
        data=docx, 
        file_name=f"Dictamen_{nombre}.docx", 
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
